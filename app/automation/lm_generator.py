"""Génère une lettre de motivation au format Word (.docx) puis la convertit en PDF."""
import subprocess
import tempfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_font(run, name: str = "Calibri", size: int = 11, bold: bool = False, color: tuple = None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_paragraph(doc: Document, text: str = "", alignment=WD_ALIGN_PARAGRAPH.LEFT,
                   space_before: int = 0, space_after: int = 6,
                   font_size: int = 11, bold: bool = False, color: tuple = None) -> None:
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        _set_font(run, size=font_size, bold=bold, color=color)
    return p


def generate_lm_docx(lm_texte: str, profil: dict, offer_title: str, offer_company: str) -> Path:
    """Génère un .docx formaté et retourne son chemin."""
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Style par défaut
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    today = date.today().strftime("%d/%m/%Y")
    nom = profil.get("nom", "")
    email = profil.get("email", "")
    telephone = profil.get("telephone", "")
    localisation = profil.get("localisation", "")

    # --- Expéditeur (gauche) ---
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(nom)
    _set_font(run, size=12, bold=True, color=(31, 73, 125))

    for line in [email, telephone, localisation]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        _set_font(run, size=10)

    doc.add_paragraph()  # espace

    # --- Date (droite) ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{localisation.split(',')[0].strip()}, le {today}")
    _set_font(run, size=11)

    doc.add_paragraph()  # espace

    # --- Destinataire ---
    _add_paragraph(doc, f"À l'attention du service recrutement", space_after=0)
    _add_paragraph(doc, offer_company, bold=True, space_after=0)
    _add_paragraph(doc, f"Objet : Candidature au poste de {offer_title}", bold=True, space_after=12)

    # --- Formule d'appel ---
    _add_paragraph(doc, "Madame, Monsieur,", space_after=12)

    # --- Corps de la lettre ---
    _skip_markers = ["veuillez agréer", "je vous prie", "dans l'attente",
                     "salutations distinguées", "sincères salutations",
                     "je reste disponible pour un entretien"]
    paragraphs = [p.strip() for p in lm_texte.strip().split("\n\n") if p.strip()]
    for para in paragraphs:
        if para.lower().startswith("madame") or para.lower().startswith("objet"):
            continue
        # Retire les lignes contenant une formule de politesse
        lines = para.split("\n")
        kept = [l for l in lines if not any(m in l.lower() for m in _skip_markers)]
        cleaned = "\n".join(kept).strip()
        if cleaned:
            _add_paragraph(doc, cleaned, space_after=8)

    doc.add_paragraph()  # espace

    # --- Formule de politesse ---
    _add_paragraph(
        doc,
        "Dans l'attente de votre retour, veuillez agréer, Madame, Monsieur, "
        "l'expression de mes salutations distinguées.",
        space_after=24,
    )

    # --- Signature ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(nom)
    _set_font(run, size=11, bold=True)

    # Sauvegarde dans un dossier persistant avec nom unique
    lm_dir = Path(__file__).parents[2] / "config" / "lettres"
    lm_dir.mkdir(exist_ok=True)
    import re as _re
    slug_company = _re.sub(r"[^a-zA-Z0-9]+", "_", offer_company)[:30].strip("_").lower()
    slug_title = _re.sub(r"[^a-zA-Z0-9]+", "_", offer_title)[:30].strip("_").lower()
    today_str = date.today().strftime("%Y%m%d")
    filename = f"lm_{slug_company}_{slug_title}_{today_str}"
    docx_path = lm_dir / f"{filename}.docx"
    doc.save(str(docx_path))
    return docx_path


def convert_docx_to_pdf(docx_path: Path) -> Path:
    """Convertit un .docx en PDF via LibreOffice headless."""
    out_dir = docx_path.parent
    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
        capture_output=True,
        text=True,
        timeout=30,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")
    pdf_path = out_dir / (docx_path.stem + ".pdf")
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF non généré : {pdf_path}")
    return pdf_path


def generate_lm_pdf(lm_texte: str, profil: dict, offer_title: str, offer_company: str) -> Path:
    """Pipeline complet : génère le .docx puis retourne le chemin du PDF."""
    docx_path = generate_lm_docx(lm_texte, profil, offer_title, offer_company)
    pdf_path = convert_docx_to_pdf(docx_path)
    return pdf_path
